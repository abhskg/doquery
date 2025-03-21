import io
import os
import tempfile
import logging
from typing import Tuple, List
from fastapi import HTTPException, status

logger = logging.getLogger(__name__)

class DocumentProcessor:
    """
    Utility class to process different document types and extract text.
    Supports text files, PDF, and DOCX documents.
    """
    
    @staticmethod
    def extract_text(content: bytes, filename: str) -> Tuple[str, str]:
        """
        Extract text from different file types based on extension.
        
        Args:
            content: The binary content of the file
            filename: The name of the file with extension
            
        Returns:
            Tuple of (extracted_text, content_type_description)
            
        Raises:
            HTTPException: If file can't be processed
        """
        # Get file extension
        _, file_ext = os.path.splitext(filename.lower())
        content_type = "text/plain"  # Default content type
        
        try:
            # Process PDF files
            if file_ext == '.pdf':
                logger.debug(f"Processing PDF file: {filename}")
                try:
                    import fitz  # PyMuPDF
                    with fitz.open(stream=content, filetype="pdf") as pdf_doc:
                        extracted_text = ""
                        for page_num in range(len(pdf_doc)):
                            page = pdf_doc[page_num]
                            extracted_text += page.get_text() + "\n\n"
                    logger.debug(f"Extracted {len(extracted_text)} characters from PDF")
                    content_type = "application/pdf"
                except ImportError:
                    logger.error("PyMuPDF is not installed. Cannot process PDF files.")
                    raise HTTPException(
                        status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                        detail="PDF processing is not available. Please install PyMuPDF library."
                    )
                except Exception as e:
                    logger.error(f"Error extracting text from PDF: {str(e)}")
                    raise HTTPException(
                        status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
                        detail=f"Could not extract text from PDF: {str(e)}"
                    )
                
            # Process DOCX files    
            elif file_ext == '.docx':
                logger.debug(f"Processing DOCX document: {filename}")
                try:
                    # Save to temp file
                    with tempfile.NamedTemporaryFile(delete=False, suffix=file_ext) as temp_file:
                        temp_file.write(content)
                        temp_path = temp_file.name
                    
                    try:
                        import docx
                        # Process DOCX using python-docx
                        doc = docx.Document(temp_path)
                        
                        # Extract all document content
                        text_parts = []
                        
                        # Extract text from paragraphs
                        for para in doc.paragraphs:
                            if para.text.strip():
                                text_parts.append(para.text)
                        
                        # Extract text from tables
                        for table in doc.tables:
                            for row in table.rows:
                                row_text = []
                                for cell in row.cells:
                                    # Get text from cell paragraphs
                                    cell_text = ' '.join([p.text for p in cell.paragraphs if p.text.strip()])
                                    if cell_text:
                                        row_text.append(cell_text)
                                if row_text:
                                    text_parts.append(' | '.join(row_text))
                        
                        # Extract text from headers
                        for section in doc.sections:
                            header = section.header
                            if header:
                                header_text = ' '.join([p.text for p in header.paragraphs if p.text.strip()])
                                if header_text:
                                    text_parts.append(f"Header: {header_text}")
                            
                            footer = section.footer
                            if footer:
                                footer_text = ' '.join([p.text for p in footer.paragraphs if p.text.strip()])
                                if footer_text:
                                    text_parts.append(f"Footer: {footer_text}")
                        
                        extracted_text = "\n".join(text_parts)
                        logger.debug(f"Extracted {len(extracted_text)} characters from DOCX")
                        content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    finally:
                        # Clean up temp file
                        os.unlink(temp_path)
                        
                except ImportError:
                    logger.error("python-docx is not installed. Cannot process DOCX files.")
                    raise HTTPException(
                        status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                        detail="DOCX processing is not available. Please install python-docx library."
                    )
                except Exception as e:
                    logger.error(f"Error extracting text from DOCX: {str(e)}")
                    raise HTTPException(
                        status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
                        detail=f"Could not extract text from DOCX: {str(e)}"
                    )
            
            # Process text files
            else:
                # For text files and other formats, try to decode as UTF-8
                try:
                    extracted_text = content.decode("utf-8")
                    logger.debug(f"Successfully decoded file content as UTF-8")
                except UnicodeDecodeError:
                    logger.warning(f"File {filename} is not UTF-8 encoded")
                    raise HTTPException(
                        status_code=status.HTTP_415_UNSUPPORTED_MEDIA_TYPE,
                        detail="File encoding not supported. Please upload a UTF-8 encoded text file, PDF, or DOCX document.",
                    )
            
            # Verify we have content
            if not extracted_text or extracted_text.strip() == "":
                logger.warning(f"No text content extracted from file {filename}")
                raise HTTPException(
                    status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
                    detail="Could not extract any text from the document. The file may be empty, corrupt, or contains only images.",
                )
                
            return extracted_text, content_type
            
        except HTTPException:
            # Re-raise HTTP exceptions
            raise
        except Exception as e:
            logger.error(f"Error processing document: {str(e)}")
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail=f"Error processing document: {str(e)}",
            ) 